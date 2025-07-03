import os
from azure.core.credentials import AzureKeyCredential
from azure.ai.documentintelligence import DocumentIntelligenceClient
from docx import Document

client = DocumentIntelligenceClient(
    endpoint="https://dai-moodle-eastus2.cognitiveservices.azure.com/", credential=AzureKeyCredential("BiqgEbh4itSZn93QSWz4OdaFZlQfFj5dT3N8YaYBh5bKUq6bUxO5JQQJ99BGACHYHv6XJ3w3AAALACOGqlZY")
)
def convert_doc_to_docx_with_spire(doc_path):
    """Convert a .doc file to .docx format using Spire.Doc."""
    try:
        from spire.doc import Document, FileFormat

        doc_path_str = str(doc_path)
        docx_path = os.path.splitext(doc_path_str)[0] + "_converted.docx"

        # Load and convert
        doc = Document()
        doc.LoadFromFile(doc_path_str)
        print(f"📄 Loaded DOC file: {doc_path_str}")

        doc.SaveToFile(docx_path, FileFormat.Docx)
        print(f"✅ Converted {doc_path_str} → {docx_path}")

        doc.Close()
        return docx_path

    except Exception as e:
        print(f"⚠️ Spire.Doc conversion failed: {e}")
        return None

def extract_text_from_docx(file_path):
    """Extract text from a DOCX file."""
    import docx

    text = ""
    try:
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"
    except Exception as e:
        text = f"Error extracting text from DOCX: {str(e)}"

    return text

def convert_text_to_markdown(text):
    """Convert plain text to markdown format."""
    if not text or text.startswith("Error"):
        return text

    # Clean up Spire.Doc warning message
    text = text.replace("Evaluation Warning: The document was created with Spire.Doc for Python.", "")

    # Simple conversion - add markdown formatting
    lines = text.split('\n')
    markdown_content = ""

    for line in lines:
        line = line.strip()
        if line:
            markdown_content += line + "\n\n"

    return markdown_content

def document_to_markdown(path: str) -> str | None:
    """
    Accepts *any* document path (.doc, .docx, .pdf, .pptx, .png, …)
    and returns a Markdown string (or None if it failed).
    """
    if not os.path.exists(path):
        print(f"❌ File not found: {path}")
        return None

    # If it's a legacy .doc, convert first
    is_doc = path.lower().endswith(".doc")
    work_path = path
    if is_doc:
        converted = convert_doc_to_docx_with_spire(path)
        if not converted:
            return None
        work_path = converted

    print(f"🔍 Processing {work_path} → MD File")
    try:
        # For converted .doc we already have the plain text – skip DI round-trip
        if is_doc:
            text = extract_text_from_docx(work_path)
            md = convert_text_to_markdown(text)
            print(f"📝 Extracted text directly from converted DOC file")
        else:
            # For other files, use Azure Document Intelligence
            with open(work_path, "rb") as f:
                poller = client.begin_analyze_document(
                    "prebuilt-layout",
                    f,
                    content_type=None,              # let the service detect type
                    output_content_format="markdown"
                )
            result = poller.result()
            md = result.content

        return md
    except Exception as e:
        print(f"❌ Error in Azure DI for {work_path}: {e}")
        return None
    finally:
        # Remove temp _converted.docx if we created one
        if is_doc and work_path.endswith("_converted.docx") and os.path.exists(work_path):
            try:
                os.remove(work_path)
            except OSError:
                pass

def save_markdown(md: str, orig_path: str, out_dir: str = "docs_md") -> str | None:
    os.makedirs(out_dir, exist_ok=True)

    base_name = os.path.splitext(os.path.basename(orig_path))[0]
    out_path = os.path.join(out_dir, f"{base_name}.md")

    try:
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(md)
        print(f"✅ Saved → {out_path}")
        return out_path
    except Exception as e:
        print(f"❌ Could not save {out_path}: {e}")
        return None


if __name__ == "__main__":
    # Example usage: replace with whatever paths you want
    files_to_process = [
        # r"Docs\sample.pdf",
        # r"Docs\DL_14_LLMs.pptx",
        r"Docs\1002806.docx",
        # r"Docs\4387683.png",
        r"Docs\4375695.doc",
    ]

    for file in files_to_process:
        md_content = document_to_markdown(file)
        if md_content:
            save_markdown(md_content, file)
